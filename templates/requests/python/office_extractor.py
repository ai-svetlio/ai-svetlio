"""
ClientRequests Office Document Extractor
=========================================
Адаптирана версия на AI_Svetlio EOOD office_processor за ClientRequests.

Извлича текст от различни видове офис документи:
- .docx (Word документи, Teams транскрипти)
- .doc (стари Word документи)
- .xlsx, .xls (Excel)
- .rtf, .txt, .xml (текстови)
- .odt (OpenDocument)
- .eml (имейли с прикачени офис документи)

Оригинал: AI_Svetlio EOOD internal office_processor (НЕ Е ПРОМЕНЯН)
Тази версия е адаптирана за ClientRequests/ системата.

Използване:
    python office_extractor.py                  # Обработва всички файлове в inbox/
    python office_extractor.py --file "X.docx"  # Обработва конкретен файл
    python office_extractor.py --file "X.eml"   # Извлича DOCX от EML и обработва

Извежда:
    processed/<filename>_extracted.json   - метаданни
    processed/<filename>_body.txt         - извлечен текст
    processed/<filename>.docx             - копие на DOCX (при EML)
"""

import json
import logging
import re
import sys
import argparse
import time
from datetime import datetime
from email import policy
from email.parser import BytesParser
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, List, Tuple
from dataclasses import dataclass

# Configure logging
LOG_FILE = Path(__file__).parent / 'office_extractor.log'
logging.basicConfig(
    level=logging.DEBUG,
    filename=str(LOG_FILE),
    filemode='a',
    format='%(asctime)s - %(levelname)s - %(message)s'
)

console = logging.StreamHandler()
console.setLevel(logging.INFO)
console.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
logging.getLogger('').addHandler(console)

logger = logging.getLogger(__name__)

# Paths - адаптирани за ClientRequests
BASE_DIR = Path(__file__).parent
INBOX_DIR = BASE_DIR / "inbox"
PROCESSED_DIR = BASE_DIR / "processed"

# Ensure directories exist
INBOX_DIR.mkdir(exist_ok=True)
PROCESSED_DIR.mkdir(exist_ok=True)


# ================== DOCUMENT TYPES ==================

class DocumentType(Enum):
    """Поддържани типове документи."""
    DOCX = ".docx"
    DOC = ".doc"
    XLSX = ".xlsx"
    XLS = ".xls"
    RTF = ".rtf"
    TXT = ".txt"
    XML = ".xml"
    ODT = ".odt"


# Optional imports - не всички са задължителни
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx не е инсталиран. DOCX файлове няма да се обработват.")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl не е инсталиран. XLSX файлове няма да се обработват.")

try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False
    logger.warning("xlrd не е инсталиран. XLS файлове няма да се обработват.")

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

# Win32com - само на Windows
try:
    import pythoncom
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False


# ================== HELPER FUNCTIONS ==================

def clean_text(text: str) -> str:
    """Почиства извлечен текст."""
    # Remove control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Normalize multiple blank lines to max 2
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    return text.strip()


def clean_text_aggressive(text: str) -> str:
    """Агресивно почистване (за офис документи без структура)."""
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    text = re.sub(r'\n\s*\n', '\n', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


# ================== TEXT EXTRACTORS ==================

class OfficeTextExtractor:
    """Извлича текст от различни видове офис документи.

    Адаптирано от: AI_Svetlio EOOD office_processor::TextExtractor
    Разлика: Без PostgreSQL, без Watchdog, с EML поддръжка.
    """

    @staticmethod
    def extract_from_docx(file_path: Path) -> Optional[str]:
        """Извлича текст от DOCX файл."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx не е наличен, DOCX не може да се обработи")
            return None
        try:
            document = DocxDocument(str(file_path))
            paragraphs = []
            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            return clean_text('\n'.join(paragraphs))
        except Exception as e:
            logger.error(f"Грешка при извличане от DOCX {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_docx_structured(file_path: Path) -> Optional[List[Dict]]:
        """Извлича структуриран текст от DOCX (параграф по параграф).
        Полезно за Teams транскрипти, където всеки параграф е реплика.
        """
        if not DOCX_AVAILABLE:
            return None
        try:
            document = DocxDocument(str(file_path))
            paragraphs = []
            for i, para in enumerate(document.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'index': i,
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            return paragraphs
        except Exception as e:
            logger.error(f"Грешка при структурирано извличане от DOCX {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_doc(file_path: Path) -> Optional[str]:
        """Извлича текст от DOC файл (стар Word формат)."""
        if not WIN32COM_AVAILABLE:
            logger.warning(f"win32com не е наличен, DOC не може да се обработи: {file_path}")
            return None
        word = None
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(str(file_path.absolute()))
            text = doc.Content.Text
            doc.Close(False)
            return clean_text_aggressive(text)
        except Exception as e:
            logger.error(f"Грешка при извличане от DOC {file_path}: {e}")
            return None
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
            try:
                pythoncom.CoUninitialize()
            except:
                pass

    @staticmethod
    def extract_from_xlsx(file_path: Path) -> Optional[str]:
        """Извлича текст от XLSX файл."""
        if not OPENPYXL_AVAILABLE:
            logger.warning(f"openpyxl не е наличен, XLSX не може да се обработи: {file_path}")
            return None
        try:
            workbook = openpyxl.load_workbook(str(file_path), data_only=True)
            text = []
            for sheet in workbook.sheetnames:
                worksheet = workbook[sheet]
                text.append(f"\n=== Лист: {sheet} ===\n")
                for row in worksheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip() and row_text.strip() != '|':
                        text.append(row_text)
            return clean_text('\n'.join(text))
        except Exception as e:
            logger.error(f"Грешка при извличане от XLSX {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_xls(file_path: Path) -> Optional[str]:
        """Извлича текст от XLS файл (стар Excel)."""
        if not XLRD_AVAILABLE:
            logger.warning(f"xlrd не е наличен, XLS не може да се обработи: {file_path}")
            return None
        try:
            workbook = xlrd.open_workbook(str(file_path))
            text = []
            for sheet_index in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sheet_index)
                text.append(f"\n=== Лист: {sheet.name} ===\n")
                for row in range(sheet.nrows):
                    row_values = sheet.row_values(row)
                    row_text = ' | '.join(str(cell) for cell in row_values if cell)
                    if row_text.strip():
                        text.append(row_text)
            return clean_text('\n'.join(text))
        except Exception as e:
            logger.error(f"Грешка при извличане от XLS {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_txt(file_path: Path) -> Optional[str]:
        """Извлича текст от TXT/RTF/XML файл."""
        try:
            # Пробваме различни кодировки
            for encoding in ['utf-8', 'cp1251', 'iso-8859-1', 'windows-1252']:
                try:
                    text = file_path.read_text(encoding=encoding)
                    return clean_text(text)
                except UnicodeDecodeError:
                    continue
            logger.warning(f"Не може да се декодира {file_path}")
            return None
        except Exception as e:
            logger.error(f"Грешка при четене на текстов файл {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_odt(file_path: Path) -> Optional[str]:
        """Извлича текст от ODT файл."""
        if not PYPANDOC_AVAILABLE:
            logger.warning(f"pypandoc не е наличен, ODT не може да се обработи: {file_path}")
            return None
        try:
            text = pypandoc.convert_file(str(file_path), 'plain')
            return clean_text(text)
        except Exception as e:
            logger.error(f"Грешка при извличане от ODT {file_path}: {e}")
            return None

    def extract(self, file_path: Path) -> Optional[str]:
        """Извлича текст от файл по разширение."""
        try:
            doc_type = DocumentType(file_path.suffix.lower())
        except ValueError:
            logger.warning(f"Неподдържан тип файл: {file_path.suffix}")
            return None

        extractors = {
            DocumentType.DOCX: self.extract_from_docx,
            DocumentType.DOC: self.extract_from_doc,
            DocumentType.XLSX: self.extract_from_xlsx,
            DocumentType.XLS: self.extract_from_xls,
            DocumentType.RTF: self.extract_from_txt,
            DocumentType.TXT: self.extract_from_txt,
            DocumentType.XML: self.extract_from_txt,
            DocumentType.ODT: self.extract_from_odt,
        }

        extractor = extractors.get(doc_type)
        if extractor:
            return extractor(file_path)
        return None


# ================== EML HANDLER ==================

class EmlDocumentExtractor:
    """Извлича офис документи от EML файлове и ги обработва."""

    OFFICE_EXTENSIONS = {ext.value for ext in DocumentType}

    @staticmethod
    def extract_attachments_from_eml(eml_path: Path) -> List[Tuple[str, bytes, str]]:
        """Извлича всички офис прикачени файлове от EML.

        Returns:
            List of (filename, data, content_type) tuples
        """
        with open(eml_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)

        attachments = []
        for part in msg.walk():
            filename = part.get_filename()
            if not filename:
                continue

            ext = Path(filename).suffix.lower()
            if ext in EmlDocumentExtractor.OFFICE_EXTENSIONS:
                data = part.get_payload(decode=True)
                if data:
                    attachments.append((
                        filename,
                        data,
                        part.get_content_type()
                    ))
                    logger.info(f"Намерен офис документ в EML: {filename} ({len(data)} bytes)")

        return attachments

    @staticmethod
    def get_eml_metadata(eml_path: Path) -> Dict:
        """Извлича метаданни от EML файл."""
        with open(eml_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)

        return {
            'from': msg.get('From', ''),
            'to': msg.get('To', ''),
            'cc': msg.get('CC', ''),
            'date': msg.get('Date', ''),
            'subject': msg.get('Subject', ''),
        }


# ================== TEAMS TRANSCRIPT PARSER ==================

class TeamsTranscriptParser:
    """Парсва Teams транскрипти от DOCX файлове.

    Teams транскриптите имат специфичен формат в DOCX:
    - Параграф 0: Заглавие (Meeting Name-YYYYMMDD_HHMMSS-Enregistrement/Recording)
    - Параграф 1: Дата (напр. "5 février 2026, 12:00PM")
    - Параграф 2: Продължителност (напр. "55min 28sec")
    - Параграф 3: Кой е стартирал транскрипцията
    - Параграфи 4+: Съобщения — всеки параграф съдържа:
      \\nSpeaker Name   M:SS
      Текст на съобщението
      (може да е на няколко реда)
    """

    # Pattern за speaker+time в началото на параграф (може да започва с \n)
    # Формат: "Speaker Name   M:SS" или "Speaker Name   HH:MM:SS"
    SPEAKER_LINE_PATTERN = re.compile(r'^(.+?)\s{2,}(\d+:\d{2}(?::\d{2})?)\s*$')

    @staticmethod
    def is_teams_transcript(paragraphs: List[Dict]) -> bool:
        """Проверява дали документът е Teams транскрипт."""
        if not paragraphs or len(paragraphs) < 5:
            return False

        # Проверяваме първите параграфи за характерни признаци
        first_texts = [p['text'] for p in paragraphs[:5]]

        indicators = 0
        for text in first_texts:
            if any(kw in text.lower() for kw in ['recording', 'enregistrement', 'запис',
                                                   'transcription', 'транскрипция',
                                                   'commencé la transcription',
                                                   'started transcription']):
                indicators += 1
            if re.search(r'\d+min|\d+мин|\d+sec|\d+сек', text):
                indicators += 1

        # Проверяваме дали параграфите съдържат speaker patterns
        # В Teams формата: параграф започва с \nSpeakerName   M:SS\nText
        speaker_count = 0
        for p in paragraphs[4:min(15, len(paragraphs))]:
            lines = p['text'].strip().split('\n')
            for line in lines:
                line = line.strip()
                if TeamsTranscriptParser.SPEAKER_LINE_PATTERN.match(line):
                    speaker_count += 1
                    break

        return indicators >= 2 or speaker_count >= 3

    @staticmethod
    def parse(paragraphs: List[Dict]) -> Dict:
        """Парсва Teams транскрипт в структуриран формат.

        В Teams DOCX формата всеки параграф (от 4-ти нататък) съдържа:
        - Ред 1: "SpeakerName   M:SS"
        - Редове 2+: Текстът на съобщението
        """
        result = {
            'title': '',
            'date': '',
            'duration': '',
            'started_by': '',
            'participants': set(),
            'messages': [],
        }

        if not paragraphs:
            return result

        # Parse header
        result['title'] = paragraphs[0]['text'].strip() if paragraphs else ''

        if len(paragraphs) > 1:
            result['date'] = paragraphs[1]['text'].strip()

        if len(paragraphs) > 2:
            result['duration'] = paragraphs[2]['text'].strip()

        if len(paragraphs) > 3:
            result['started_by'] = paragraphs[3]['text'].strip()

        # Parse messages — всеки параграф е отделно съобщение
        for p in paragraphs[4:]:
            raw_text = p['text']
            lines = raw_text.split('\n')

            # Намери speaker line
            speaker = None
            time_stamp = None
            text_lines = []

            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                if not speaker:
                    # Опитваме да match-нем speaker+time
                    match = TeamsTranscriptParser.SPEAKER_LINE_PATTERN.match(line_stripped)
                    if match:
                        speaker = match.group(1).strip()
                        time_stamp = match.group(2).strip()
                        continue

                    # Ако не е speaker line, може да е продължение от header
                    # или специален ред — пропускаме
                    continue

                # Всичко след speaker line е текст
                if line_stripped:
                    text_lines.append(line_stripped)

            if speaker and text_lines:
                result['messages'].append({
                    'speaker': speaker,
                    'time': time_stamp or '',
                    'text': '\n'.join(text_lines)
                })
                result['participants'].add(speaker)
            elif speaker:
                # Speaker без текст (напр. "arrêt de la transcription")
                # Не добавяме като съобщение, но записваме участника
                result['participants'].add(speaker)

        # Convert set to sorted list
        result['participants'] = sorted(result['participants'])

        result['summary_stats'] = {
            'total_messages': len(result['messages']),
            'participants_count': len(result['participants']),
            'participants': result['participants'],
            'duration': result['duration'],
        }

        return result


# ================== MAIN PROCESSOR ==================

class OfficeDocumentProcessor:
    """Обработва офис документи за ClientRequests системата."""

    SUPPORTED_EXTENSIONS = {ext.value for ext in DocumentType} | {'.eml'}

    def __init__(self):
        self.extractor = OfficeTextExtractor()
        self.processed_count = 0

    def is_supported(self, file_path: Path) -> bool:
        """Проверява дали файлът е поддържан."""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def process_file(self, file_path: Path) -> Optional[Dict]:
        """Обработва един файл и връща резултат."""
        logger.info(f"Обработка на файл: {file_path}")

        ext = file_path.suffix.lower()

        if ext == '.eml':
            return self._process_eml(file_path)
        elif ext in {e.value for e in DocumentType}:
            return self._process_office_doc(file_path)
        else:
            logger.warning(f"Неподдържан формат: {ext}")
            return None

    def _process_eml(self, eml_path: Path) -> Optional[Dict]:
        """Обработва EML файл — извлича прикачени офис документи."""
        logger.info(f"Обработка на EML с офис документи: {eml_path}")

        # Get EML metadata
        eml_meta = EmlDocumentExtractor.get_eml_metadata(eml_path)

        # Extract office attachments
        attachments = EmlDocumentExtractor.extract_attachments_from_eml(eml_path)

        if not attachments:
            logger.info(f"Няма офис документи в EML: {eml_path}")
            return None

        results = []
        for filename, data, content_type in attachments:
            # Save attachment to processed/
            att_path = PROCESSED_DIR / filename
            if att_path.exists():
                stem = att_path.stem
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                att_path = PROCESSED_DIR / f"{stem}_{ts}{att_path.suffix}"

            att_path.write_bytes(data)
            logger.info(f"Записан прикачен файл: {att_path} ({len(data)} bytes)")

            # Process the extracted file
            result = self._process_office_doc(att_path, eml_metadata=eml_meta)
            if result:
                result['source_eml'] = eml_path.name
                results.append(result)

        if results:
            # Move EML to processed
            dest = PROCESSED_DIR / eml_path.name
            if dest.exists():
                stem = eml_path.stem
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                dest = PROCESSED_DIR / f"{stem}_{ts}{eml_path.suffix}"
            eml_path.rename(dest)
            logger.info(f"Преместен EML: {eml_path} -> {dest}")

            self.processed_count += 1
            # Return first/main result
            return results[0] if len(results) == 1 else {
                'source_file': eml_path.name,
                'format': 'eml_multi',
                'documents': results,
                'eml_metadata': eml_meta
            }

        return None

    def _process_office_doc(self, file_path: Path, eml_metadata: Dict = None) -> Optional[Dict]:
        """Обработва офис документ."""
        logger.info(f"Обработка на офис документ: {file_path}")

        result = {
            'source_file': file_path.name,
            'format': file_path.suffix.lower().lstrip('.'),
            'file_size_kb': file_path.stat().st_size / 1024,
            'processing_timestamp': datetime.now().isoformat(),
            'extracted_text': '',
            'is_teams_transcript': False,
            'teams_data': None,
        }

        if eml_metadata:
            result['eml_metadata'] = eml_metadata

        # Try structured extraction first for DOCX (Teams transcripts)
        if file_path.suffix.lower() == '.docx':
            paragraphs = OfficeTextExtractor.extract_from_docx_structured(file_path)

            if paragraphs and TeamsTranscriptParser.is_teams_transcript(paragraphs):
                logger.info(f"Разпознат Teams транскрипт: {file_path}")
                result['is_teams_transcript'] = True
                teams_data = TeamsTranscriptParser.parse(paragraphs)
                result['teams_data'] = teams_data

                # Build readable text from transcript
                lines = []
                lines.append(f"# {teams_data['title']}")
                lines.append(f"Дата: {teams_data['date']}")
                lines.append(f"Продължителност: {teams_data['duration']}")
                lines.append(f"Участници: {', '.join(teams_data['participants'])}")
                lines.append(f"{teams_data.get('started_by', '')}")
                lines.append("")
                lines.append("---")
                lines.append("")

                for msg in teams_data['messages']:
                    lines.append(f"**{msg['speaker']}** [{msg['time']}]:")
                    lines.append(msg['text'])
                    lines.append("")

                result['extracted_text'] = '\n'.join(lines)
            else:
                # Normal DOCX
                result['extracted_text'] = self.extractor.extract(file_path) or ''
        else:
            # Other formats
            result['extracted_text'] = self.extractor.extract(file_path) or ''

        if not result['extracted_text']:
            logger.warning(f"Не е извлечен текст от: {file_path}")
            return result

        # Save extracted data
        self._save_results(file_path, result)

        return result

    def _save_results(self, file_path: Path, result: Dict) -> None:
        """Запазва резултатите."""
        stem = file_path.stem

        # Save JSON metadata
        json_path = PROCESSED_DIR / f"{stem}_extracted.json"
        json_data = {k: v for k, v in result.items() if k != 'extracted_text'}

        # Serialize teams_data participants set
        if json_data.get('teams_data') and json_data['teams_data'].get('participants'):
            json_data['teams_data']['participants'] = list(json_data['teams_data']['participants'])

        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2, default=str)
        logger.info(f"Записани метаданни: {json_path}")

        # Save body text
        body_path = PROCESSED_DIR / f"{stem}_body.txt"

        header_lines = []
        if result.get('eml_metadata'):
            meta = result['eml_metadata']
            header_lines.append(f"От: {meta.get('from', '')}")
            header_lines.append(f"До: {meta.get('to', '')}")
            if meta.get('cc'):
                header_lines.append(f"CC: {meta['cc']}")
            header_lines.append(f"Дата: {meta.get('date', '')}")
            header_lines.append(f"Тема: {meta.get('subject', '')}")

        header_lines.append(f"Файл: {result.get('source_file', '')}")
        header_lines.append(f"Формат: {result.get('format', '')}")

        if result.get('is_teams_transcript') and result.get('teams_data'):
            td = result['teams_data']
            header_lines.append(f"Тип: Teams транскрипт")
            header_lines.append(f"Участници: {', '.join(td.get('participants', []))}")
            header_lines.append(f"Продължителност: {td.get('duration', '')}")
            header_lines.append(f"Съобщения: {td['summary_stats']['total_messages']}")

        header_lines.append(f"Размер: {result.get('file_size_kb', 0):.1f} KB")

        with open(body_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(header_lines))
            f.write(f"\n\n{'='*60}\n\n")
            f.write(result.get('extracted_text', ''))

        logger.info(f"Записано тяло: {body_path}")

    def process_inbox(self) -> List[Dict]:
        """Обработва всички поддържани файлове в inbox/."""
        results = []

        files = sorted(INBOX_DIR.iterdir())
        if not files:
            logger.info("inbox/ е празна.")
            return results

        for file_path in files:
            if file_path.is_file() and self.is_supported(file_path):
                result = self.process_file(file_path)
                if result:
                    results.append(result)

        logger.info(f"Обработени {len(results)} файла")
        return results


# ================== CLI ==================

def print_result(result: Dict):
    """Показва резюме на резултата."""
    print(f"\n{'='*60}")
    print(f"  Файл:     {result.get('source_file', '')}")
    print(f"  Формат:   {result.get('format', '')}")
    print(f"  Размер:   {result.get('file_size_kb', 0):.1f} KB")

    if result.get('eml_metadata'):
        meta = result['eml_metadata']
        print(f"  От:       {meta.get('from', '')}")
        print(f"  До:       {meta.get('to', '')}")
        print(f"  Тема:     {meta.get('subject', '')}")

    if result.get('is_teams_transcript') and result.get('teams_data'):
        td = result['teams_data']
        print(f"  Тип:      Teams транскрипт")
        print(f"  Участ.:   {', '.join(td.get('participants', []))}")
        print(f"  Продълж.: {td.get('duration', '')}")
        print(f"  Съобщ.:   {td['summary_stats']['total_messages']}")

    text = result.get('extracted_text', '')
    print(f"  Текст:    {len(text)} символа")
    print(f"{'='*60}\n")


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    parser = argparse.ArgumentParser(description='ClientRequests Office Document Extractor')
    parser.add_argument('--file', type=str, help='Обработи конкретен файл')
    args = parser.parse_args()

    processor = OfficeDocumentProcessor()

    if args.file:
        file_path = Path(args.file)
        if not file_path.exists():
            file_path = INBOX_DIR / args.file
        if not file_path.exists():
            print(f"Файлът не съществува: {args.file}")
            sys.exit(1)

        result = processor.process_file(file_path)
        if result:
            print_result(result)
        else:
            print("Грешка или няма извлечен текст.")
            sys.exit(1)
    else:
        results = processor.process_inbox()
        if results:
            for result in results:
                print_result(result)
            print(f"Общо обработени: {len(results)} файла")
        else:
            print("Няма поддържани файлове в inbox/")


if __name__ == '__main__':
    main()
